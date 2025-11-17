# =============================================================================
# PDF Toolkit
# Version: 0.1.0
# Description: Command-line utility toolkit for common PDF manipulations.
# =============================================================================

from __future__ import annotations

import argparse
import io
import json
import sys
from contextlib import ExitStack
from pathlib import Path
from typing import Any, Dict, List, Sequence

try:
    import fitz  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    fitz = None

try:
    from tqdm import tqdm
except ImportError:  # pragma: no cover - handled at runtime
    def tqdm(iterable, **_: object):
        return iterable

try:
    import pikepdf  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    pikepdf = None

try:
    from PIL import Image  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    Image = None

try:
    import pytesseract  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    pytesseract = None

try:
    from docx import Document  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    Document = None

try:
    from odf.opendocument import OpenDocumentText  # type: ignore[import-not-found]
    from odf.text import P  # type: ignore[import-not-found]
    from odf import style  # type: ignore[import-not-found]
    from odf.style import Style, TextProperties, ParagraphProperties  # type: ignore[import-not-found]
    odf_available = True
except ImportError:  # pragma: no cover - handled at runtime
    odf_available = False
    OpenDocumentText = None
    P = None
    style = None
    Style = None
    TextProperties = None
    ParagraphProperties = None


__version__ = "0.3.0"


class _NullStream:
    """Fallback stream used when stdout/stderr are unavailable (e.g., PyInstaller GUI)."""

    def write(self, *_: object, **__: object) -> None:
        pass

    def flush(self) -> None:
        pass


if getattr(sys, "stdout", None) is None:
    sys.stdout = _NullStream()  # type: ignore[assignment]

if getattr(sys, "stderr", None) is None:
    sys.stderr = _NullStream()  # type: ignore[assignment]

# ============= 工具函數區 =============

def parse_page_spec(spec: str, total_pages: int) -> list[int]:
    """
    Convert a 1-based page specification string into a sorted list of 0-based indexes.

    Args:
        spec: Page specification string, e.g. "1,3,5-7,10-".
        total_pages: Total number of pages available in the PDF.

    Returns:
        A sorted list of unique 0-based page indexes extracted from the specification.

    Raises:
        ValueError: If the specification has invalid syntax or references out-of-range pages.
    """
    if total_pages <= 0:
        raise ValueError("Total pages must be a positive integer.")

    stripped_spec = spec.strip()
    if not stripped_spec:
        return []

    indexes: set[int] = set()

    for raw_token in stripped_spec.split(","):
        token = raw_token.strip()
        if not token:
            raise ValueError("頁碼範圍中存在空白項，請檢查輸入。")

        if token.count("-") > 1:
            raise ValueError(f"頁碼範圍格式錯誤：'{token}'。")

        if "-" in token:
            start_str, end_str = token.split("-")
            start = _parse_positive_int(start_str, "起始頁碼") if start_str else None
            end = _parse_positive_int(end_str, "結束頁碼") if end_str else None

            if start is None and end is None:
                raise ValueError(f"頁碼範圍格式錯誤：'{token}'。")

            start_index = start - 1 if start is not None else 0
            end_index = (end - 1) if end is not None else total_pages - 1

            if start_index < 0 or start_index >= total_pages:
                raise ValueError(f"起始頁碼超出範圍：{start}.")
            if end_index < 0 or end_index >= total_pages:
                raise ValueError(f"結束頁碼超出範圍：{end}.")
            if start_index > end_index:
                raise ValueError(f"頁碼範圍需由小到大：'{token}'。")

            indexes.update(range(start_index, end_index + 1))
            continue

        page_number = _parse_positive_int(token, "頁碼")
        page_index = page_number - 1
        if page_index < 0 or page_index >= total_pages:
            raise ValueError(f"頁碼超出範圍：{page_number}.")
        indexes.add(page_index)

    return sorted(indexes)


def check_file_exists(filepath: str) -> None:
    """
    Ensure the given file path exists.

    Args:
        filepath: Path to the file that should exist.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    path = Path(filepath)
    if not path.is_file():
        raise FileNotFoundError(f"找不到檔案：{path.resolve()}")


def safe_open_pdf(filepath: str, password: str | None = None):
    """
    Safely open a PDF document using PyMuPDF.

    Args:
        filepath: Path to the PDF file.
        password: Optional password for encrypted documents.

    Returns:
        An instance of fitz.Document representing the opened PDF.

    Raises:
        ImportError: If PyMuPDF (fitz) is not installed.
        FileNotFoundError: If the specified file cannot be found.
        PermissionError: If the PDF is encrypted and no valid password is provided.
        ValueError: If the file is not a valid or readable PDF.
    """
    if fitz is None:
        raise ImportError(
            "PyMuPDF (fitz) 尚未安裝，請先執行 'pip install PyMuPDF>=1.23.0' 後再試。"
        )

    check_file_exists(filepath)

    try:
        # Deferred import keeps static analyzers aware of the module attribute.
        document = fitz.open(filepath)

        # Handle encrypted PDFs
        if document.is_encrypted:
            if password:
                # Try to authenticate with provided password
                if not document.authenticate(password):
                    document.close()
                    raise PermissionError("此 PDF 已加密，提供的密碼不正確。")
            else:
                # No password provided for encrypted PDF
                document.close()
                raise PermissionError("此 PDF 已加密，請提供正確的密碼後再試。")
    except PermissionError:
        # Re-raise permission errors as-is
        raise
    except RuntimeError as exc:
        error_message = str(exc).lower()
        if "password" in error_message or "encryption" in error_message:
            raise PermissionError("此 PDF 已加密，請提供正確的密碼後再試。") from exc
        raise ValueError(f"無法讀取 PDF 檔案：{filepath}") from exc
    except Exception as exc:  # pragma: no cover - PyMuPDF specific errors
        raise ValueError(f"無法開啟 PDF 檔案：{filepath}") from exc

    return document


def _ensure_fitz_available() -> None:
    """Raise an informative error when PyMuPDF is unavailable."""

    if fitz is None:
        raise ImportError(
            "PyMuPDF (fitz) 尚未安裝，請先執行 'pip install PyMuPDF>=1.23.0' 後再試。"
        )


def _field_type_name(field_type: int) -> str:
    """Return a human-friendly widget type name."""

    if fitz is None:
        return "未知"

    type_names = {
        fitz.PDF_WIDGET_TYPE_TEXT: "文字",
        fitz.PDF_WIDGET_TYPE_CHECKBOX: "核取方塊",
        fitz.PDF_WIDGET_TYPE_COMBOBOX: "下拉選單",
        fitz.PDF_WIDGET_TYPE_LISTBOX: "列表",
        fitz.PDF_WIDGET_TYPE_RADIOBUTTON: "單選按鈕",
        fitz.PDF_WIDGET_TYPE_BUTTON: "按鈕",
        fitz.PDF_WIDGET_TYPE_SIGNATURE: "簽名",
    }

    return type_names.get(field_type, f"未知類型 ({field_type})")


def _normalize_checkbox_value(widget, value: Any) -> str:
    """Return a proper checkbox state string based on the provided value."""

    truthy = {"1", "true", "yes", "y", "on", "checked"}

    should_check: bool
    if isinstance(value, bool):
        should_check = value
    elif isinstance(value, (int, float)):
        should_check = value != 0
    elif isinstance(value, str):
        should_check = value.strip().lower() in truthy
    else:
        should_check = bool(value)

    on_state = getattr(widget, "button_on_state", None) or getattr(
        widget, "on_state_name", None
    )

    if not on_state:
        on_state = "Yes"

    return on_state if should_check else "Off"


def _load_json_data(filepath: str | None) -> dict[str, Any]:
    """Load key/value mappings from a JSON file if provided."""

    if not filepath:
        return {}

    path = Path(filepath)
    if not path.is_file():
        raise FileNotFoundError(f"找不到資料檔案：{path.resolve()}")

    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON 解析失敗：{filepath}") from exc

    if not isinstance(data, dict):
        raise ValueError("資料檔案內容必須為物件（鍵值對）。")

    return data


def _parse_key_value_pairs(pairs: Sequence[str] | None) -> dict[str, Any]:
    """Parse a list of KEY=VALUE strings into a dictionary."""

    if not pairs:
        return {}

    parsed: dict[str, Any] = {}
    for pair in pairs:
        if "=" not in pair:
            raise ValueError(f"資料參數必須為 key=value 格式：'{pair}'")
        key, value = pair.split("=", 1)
        key = key.strip()
        if not key:
            raise ValueError(f"資料鍵不得為空值：'{pair}'")
        parsed[key] = value.strip()

    return parsed


def extract_pdf_form_fields(pdf_path: str) -> list[dict[str, Any]]:
    """Return metadata for all interactive form fields inside a PDF document."""

    _ensure_fitz_available()

    with ExitStack() as stack:
        doc = safe_open_pdf(pdf_path)
        stack.callback(doc.close)
        fields: list[dict[str, Any]] = []

        for page_index, page in enumerate(doc, start=1):
            widgets = page.widgets() or []
            for widget in widgets:
                name = widget.field_name or ""
                if not name:
                    continue

                field_info: dict[str, Any] = {
                    "name": name,
                    "type": _field_type_name(widget.field_type),
                    "page": page_index,
                    "value": widget.field_value,
                }

                options = getattr(widget, "choice_values", None)
                if options:
                    field_info["options"] = list(options)

                fields.append(field_info)

        return fields


def fill_pdf_form(
    template_path: str,
    output_path: str,
    data: Dict[str, Any],
    *,
    flatten: bool = False,
) -> dict[str, Any]:
    """Fill a PDF form using provided data and save the result."""

    _ensure_fitz_available()

    if not data:
        raise ValueError("未提供任何填寫資料。請使用 --data 或 --value 指定欄位內容。")

    output_file = Path(output_path)
    if output_file.exists() and output_file.is_dir():
        raise IsADirectoryError(f"輸出路徑為資料夾：{output_file}")

    filled_fields: dict[str, Any] = {}

    with ExitStack() as stack:
        doc = safe_open_pdf(template_path)
        stack.callback(doc.close)

        for page in doc:
            widgets = page.widgets() or []
            for widget in widgets:
                field_name = widget.field_name
                if not field_name or field_name not in data:
                    continue

                value = data[field_name]
                try:
                    if widget.field_type == fitz.PDF_WIDGET_TYPE_TEXT:
                        widget.field_value = "" if value is None else str(value)
                    elif widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                        widget.field_value = _normalize_checkbox_value(widget, value)
                    elif widget.field_type in (
                        fitz.PDF_WIDGET_TYPE_COMBOBOX,
                        fitz.PDF_WIDGET_TYPE_LISTBOX,
                    ):
                        widget.field_value = "" if value is None else str(value)
                    elif widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                        widget.field_value = "" if value is None else str(value)
                    else:
                        widget.field_value = "" if value is None else str(value)

                    widget.update()
                    filled_fields[field_name] = value
                except Exception as exc:
                    raise ValueError(
                        f"欄位 '{field_name}' 無法填寫，請確認提供的資料是否符合欄位型別。"
                    ) from exc

        if not filled_fields:
            raise ValueError("提供的資料沒有對應到任何表單欄位。")

        save_kwargs: dict[str, Any] = {"deflate": 1, "incremental": 0}
        if flatten:
            save_kwargs["appearance"] = 1
            save_kwargs["clean"] = 1

        output_parent = output_file.parent
        if output_parent and not output_parent.exists():
            output_parent.mkdir(parents=True, exist_ok=True)

        try:
            doc.save(output_file, **save_kwargs)
        except OSError as exc:  # pragma: no cover - depends on filesystem
            raise OSError(f"無法寫入輸出檔案：{output_file}") from exc

    return filled_fields


def format_file_size(size_bytes: int) -> str:
    """
    Format a file size (in bytes) into a human-readable string using binary units.

    Args:
        size_bytes: File size in bytes.

    Returns:
        A human-readable string representing the size (e.g., "1.5 KB").
    """
    if size_bytes < 0:
        raise ValueError("檔案大小必須為非負整數。")

    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    value = float(size_bytes)

    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f} {unit}"
        value /= 1024

    # Fallback - should not reach here due to loop return.
    return f"{size_bytes} B"


def _parse_positive_int(value: str, label: str) -> int:
    """
    Parse a positive integer helper used by parse_page_spec.

    Args:
        value: Raw string to parse.
        label: Human readable label for error messages.

    Returns:
        Parsed positive integer.

    Raises:
        ValueError: If parsing fails or the number is not positive.
    """
    if not value:
        raise ValueError(f"{label}不可為空。")
    if not value.isdigit():
        raise ValueError(f"{label}必須為正整數：'{value}'。")
    number = int(value)
    if number <= 0:
        raise ValueError(f"{label}必須大於零：{number}。")
    return number


# ============= 基礎 PDF 操作區 =============

def merge_pdfs(input_pdfs: Sequence[str], output_pdf: str) -> None:
    """
    Merge multiple PDF files into a single document.

    Args:
        input_pdfs: Ordered collection of PDF paths to merge.
        output_pdf: Destination PDF path.

    Raises:
        ImportError: If PyMuPDF is not installed.
        ValueError: If no input files are provided or a source is encrypted.
        FileNotFoundError: If any input file is missing.
        OSError: If writing the output fails.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法執行合併功能。")

    if not input_pdfs:
        raise ValueError("請至少提供一個要合併的 PDF 檔案。")

    print(f"合併 {len(input_pdfs)} 個檔案...")

    with ExitStack() as stack:
        source_documents: List["fitz.Document"] = []
        for path in input_pdfs:
            try:
                document = safe_open_pdf(path)
            except PermissionError as exc:
                raise ValueError(f"檔案已加密，無法合併：{path}") from exc
            stack.callback(document.close)
            source_documents.append(document)

        output_document = fitz.open()
        stack.callback(output_document.close)

        total_pages = 0
        for document in tqdm(
            source_documents,
            desc="合併 PDF",
            unit="檔",
        ):
            output_document.insert_pdf(document)
            total_pages += document.page_count

        try:
            output_document.save(output_pdf)
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc

    print(f"✓ 成功合併 {len(input_pdfs)} 個檔案，總共 {total_pages} 頁")


def split_pdf(input_pdf: str, output_dir: str, page_spec: str | None = None) -> None:
    """
    Split a PDF into separate documents by single pages or ranges.

    Args:
        input_pdf: Source PDF path.
        output_dir: Directory where split files will be created.
        page_spec: Optional page specification string for selective splitting.

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the page specification is invalid or empty.
        OSError: If the output directory cannot be created or written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法執行拆分功能。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法拆分：{input_pdf}") from exc

    try:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        document.close()
        raise OSError(f"無法建立輸出資料夾：{output_dir}") from exc

    base_name = Path(input_pdf).stem
    total_pages = document.page_count
    print(f"拆分 PDF（總共 {total_pages} 頁）...")

    try:
        if page_spec is None:
            groups = [[page_index] for page_index in range(total_pages)]
        else:
            indexes = parse_page_spec(page_spec, total_pages)
            if not indexes:
                raise ValueError("頁碼範圍解析結果為空，請確認輸入。")
            groups = _group_consecutive(indexes)

        created_files = 0
        for group in tqdm(groups, desc="拆分 PDF", unit="檔"):
            start_index = group[0]
            end_index = group[-1]
            new_document = fitz.open()
            for page_index in group:
                new_document.insert_pdf(
                    document,
                    from_page=page_index,
                    to_page=page_index,
                )

            if len(group) == 1:
                filename = f"{base_name}_page_{start_index + 1:03d}.pdf"
            else:
                filename = (
                    f"{base_name}_pages_{start_index + 1:03d}-{end_index + 1:03d}.pdf"
                )

            output_file = output_path / filename
            try:
                new_document.save(output_file.as_posix())
            except OSError as exc:
                new_document.close()
                raise OSError(f"無法寫入輸出檔案：{output_file}") from exc
            finally:
                new_document.close()

            created_files += 1

    finally:
        document.close()

    print(f"✓ 成功拆分為 {created_files} 個檔案")


def _group_consecutive(indexes: Sequence[int]) -> List[List[int]]:
    """
    Group a sorted sequence of indexes into consecutive ranges.

    Args:
        indexes: Sorted sequence of integers.

    Returns:
        A list containing grouped consecutive index lists.
    """
    if not indexes:
        return []

    groups: List[List[int]] = []
    current_group: List[int] = [indexes[0]]

    for index in indexes[1:]:
        if index == current_group[-1] + 1:
            current_group.append(index)
            continue
        groups.append(current_group)
        current_group = [index]

    groups.append(current_group)
    return groups


def delete_pages(input_pdf: str, output_pdf: str, page_spec: str) -> None:
    """
    Delete specified pages from a PDF and write the result to a new file.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        page_spec: Page specification string indicating pages to delete.

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the page specification is invalid or empty.
        OSError: If the output file cannot be written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法執行刪除功能。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法刪除頁面：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        page_indexes = parse_page_spec(page_spec, total_pages)
        if not page_indexes:
            raise ValueError("請提供至少一個要刪除的頁碼。")

        print(f"刪除 {len(page_indexes)} 個頁面...")
        for page_index in sorted(page_indexes, reverse=True):
            document.delete_page(page_index)

        remaining_pages = document.page_count
        if remaining_pages == 0:
            print("⚠ 已刪除所有頁面，輸出檔案將為空。")

        save_kwargs = {
            "garbage": 4,  # remove unused objects so file size reflects deletions
            "deflate": True,
            "clean": True,
            "incremental": False,
        }

        try:
            document.save(output_pdf, **save_kwargs)
        except TypeError:
            # Older PyMuPDF versions may not support all save args.
            try:
                document.save(output_pdf, garbage=4)
            except OSError as exc:
                raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc
    finally:
        document.close()

    print(
        f"✓ 成功刪除 {len(page_indexes)} 個頁面（原 {total_pages} 頁 → 剩餘 {remaining_pages} 頁）"
    )


def rotate_pages(input_pdf: str, output_pdf: str, page_spec: str, angle: int) -> None:
    """
    Rotate selected pages in a PDF document.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        page_spec: Page specification containing pages to rotate.
        angle: Rotation angle (must be 90, 180, or 270).

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the page specification is invalid or the angle unsupported.
        OSError: If the output file cannot be written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法執行旋轉功能。")
    if angle not in {90, 180, 270}:
        raise ValueError("旋轉角度僅支援 90、180、270 度。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法旋轉頁面：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        page_indexes = parse_page_spec(page_spec, total_pages)
        if not page_indexes:
            raise ValueError("請提供至少一個要旋轉的頁碼。")

        print(f"旋轉 {len(page_indexes)} 個頁面 {angle} 度...")
        for page_index in tqdm(page_indexes, desc="旋轉頁面", unit="頁"):
            page = document[page_index]
            new_rotation = (page.rotation + angle) % 360
            page.set_rotation(new_rotation)

        try:
            document.save(output_pdf)
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc
    finally:
        document.close()

    print(f"✓ 成功旋轉 {len(page_indexes)} 個頁面")


# ============= 內容編輯區 =============

def add_watermark(
    input_pdf: str,
    output_pdf: str,
    text: str,
    size: int = 36,
    alpha: float = 0.3,
    angle: int = 0,
    color: tuple[float, float, float] = (0.5, 0.5, 0.5),
) -> None:
    """
    Add a text watermark across every page of the PDF.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        text: Watermark text content.
        size: Watermark font size (default 36).
        alpha: Opacity between 0 and 1 (default 0.3).
        angle: Rotation angle applied to the watermark - must be 0, 90, 180, or 270 (default 0).
        color: Text color expressed as an RGB tuple with values in [0, 1].

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If parameters are invalid.
        OSError: If the output file cannot be written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法添加水印。")
    if not text:
        raise ValueError("水印文字不可為空。")
    if not (0 < alpha <= 1):
        raise ValueError("水印透明度 alpha 必須介於 0 與 1 之間。")
    if angle not in {0, 90, 180, 270}:
        raise ValueError("水印旋轉角度僅支援 0、90、180、270 度。")
    if len(color) != 3 or any(component < 0 or component > 1 for component in color):
        raise ValueError("顏色需為三個 0-1 之間的浮點值，例如 (0.5, 0.5, 0.5)。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法添加水印：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        print(f"添加水印 \"{text}\" 到 {total_pages} 頁...")
        for page_index in tqdm(range(total_pages), desc="添加水印", unit="頁"):
            page = document[page_index]
            text_rect = fitz.Rect(0, 0, page.rect.width, page.rect.height)
            page.insert_textbox(
                text_rect,
                text,
                fontsize=size,
                rotate=angle,
                color=color,
                align=fitz.TEXT_ALIGN_CENTER,
                overlay=True,
                fill_opacity=alpha,
            )
        try:
            document.save(output_pdf)
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc
    finally:
        document.close()

    print(f"✓ 成功添加水印到 {total_pages} 頁")


# ============= 資訊查詢區 =============

def get_pdf_info(input_pdf: str) -> dict[str, str | int | bool]:
    """
    Collect metadata and basic information from a PDF file.

    Args:
        input_pdf: Source PDF path.

    Returns:
        Dictionary containing PDF details.

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the PDF cannot be read.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法查詢 PDF 資訊。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法讀取資訊：{input_pdf}") from exc

    try:
        file_path = Path(input_pdf)
        file_size = file_path.stat().st_size
        metadata = document.metadata or {}

        info: dict[str, str | int | bool] = {
            "filename": file_path.name,
            "file_size": format_file_size(file_size),
            "page_count": document.page_count,
            "is_encrypted": document.is_encrypted,
            "pdf_version": metadata.get("format") or "N/A",
            "title": metadata.get("title") or "N/A",
            "author": metadata.get("author") or "N/A",
            "subject": metadata.get("subject") or "N/A",
            "creator": metadata.get("creator") or "N/A",
            "producer": metadata.get("producer") or "N/A",
            "creation_date": metadata.get("creationDate") or "N/A",
            "mod_date": metadata.get("modDate") or "N/A",
        }
    finally:
        document.close()

    return info


def print_pdf_info(input_pdf: str) -> None:
    """
    Print PDF metadata in a user-friendly format.

    Args:
        input_pdf: Source PDF path.

    Raises:
        Same as get_pdf_info.
    """
    info = get_pdf_info(input_pdf)
    print("PDF 資訊")
    print("========")
    print(f"檔案名稱: {info['filename']}")
    print(f"檔案大小: {info['file_size']}")
    print(f"頁數: {info['page_count']} 頁")
    print(f"是否加密: {'是' if info['is_encrypted'] else '否'}")
    print(f"PDF 版本: {info['pdf_version']}")
    print(f"標題: {info['title']}")
    print(f"作者: {info['author']}")
    print(f"主題: {info['subject']}")
    print(f"建立程式: {info['creator']}")
    print(f"製作程式: {info['producer']}")
    print(f"建立日期: {info['creation_date']}")
    print(f"修改日期: {info['mod_date']}")


# ============= 文字提取區 =============

def extract_text(input_pdf: str, output_txt: str | None = None, page_spec: str | None = None) -> str:
    """
    Extract text content from a PDF file.

    Args:
        input_pdf: Source PDF path.
        output_txt: Optional output text file path.
        page_spec: Optional page specification string for selective extraction.

    Returns:
        Extracted text content.

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法提取文字。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法提取文字：{input_pdf}") from exc

    try:
        total_pages = document.page_count

        if page_spec is None:
            page_indexes = list(range(total_pages))
        else:
            page_indexes = parse_page_spec(page_spec, total_pages)
            if not page_indexes:
                raise ValueError("頁碼範圍解析結果為空，請確認輸入。")

        extracted_text = []
        print(f"提取文字（共 {len(page_indexes)} 頁）...")

        for page_index in tqdm(page_indexes, desc="提取文字", unit="頁"):
            page = document[page_index]
            text = page.get_text()
            extracted_text.append(f"--- 第 {page_index + 1} 頁 ---\n{text}\n")

        result = "\n".join(extracted_text)

        if output_txt:
            output_path = Path(output_txt)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(result, encoding="utf-8")
            print(f"✓ 文字已保存到：{output_path.resolve()}")
        else:
            print(result)

        return result
    finally:
        document.close()


# ============= 加密解密區 =============

def encrypt_pdf(input_pdf: str, output_pdf: str, user_password: str, owner_password: str | None = None) -> None:
    """
    Encrypt a PDF file with password protection.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        user_password: Password required to open the PDF.
        owner_password: Optional password for owner permissions (defaults to user_password).

    Raises:
        ImportError: If pikepdf is not installed.
        FileNotFoundError: If the input PDF does not exist.
        ValueError: If passwords are invalid.
        OSError: If the output file cannot be written.
    """
    if pikepdf is None:
        raise ImportError("pikepdf 尚未安裝，無法執行加密功能。")

    if not user_password:
        raise ValueError("使用者密碼不可為空。")

    source_path = Path(input_pdf)
    if not source_path.is_file():
        raise FileNotFoundError(f"找不到檔案：{source_path.resolve()}")

    owner_pwd = owner_password or user_password

    print(f"加密 PDF...")
    try:
        with pikepdf.open(input_pdf) as pdf:
            pdf.save(
                output_pdf,
                encryption=pikepdf.Encryption(
                    user=user_password,
                    owner=owner_pwd,
                    R=6,  # AES-256 encryption
                )
            )
    except OSError as exc:
        raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc

    print(f"✓ PDF 已加密並保存到：{Path(output_pdf).resolve()}")


def decrypt_pdf(input_pdf: str, output_pdf: str, password: str) -> None:
    """
    Decrypt a password-protected PDF file.

    Args:
        input_pdf: Source encrypted PDF path.
        output_pdf: Destination unencrypted PDF path.
        password: Password to decrypt the PDF.

    Raises:
        ImportError: If pikepdf is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the password is incorrect.
        OSError: If the output file cannot be written.
    """
    if pikepdf is None:
        raise ImportError("pikepdf 尚未安裝，無法執行解密功能。")

    source_path = Path(input_pdf)
    if not source_path.is_file():
        raise FileNotFoundError(f"找不到檔案：{source_path.resolve()}")

    print(f"解密 PDF...")
    try:
        with pikepdf.open(input_pdf, password=password) as pdf:
            pdf.save(output_pdf)
    except pikepdf.PasswordError as exc:  # type: ignore[attr-defined]
        raise PermissionError(f"密碼錯誤，無法解密：{input_pdf}") from exc
    except OSError as exc:
        raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc

    print(f"✓ PDF 已解密並保存到：{Path(output_pdf).resolve()}")


# ============= 圖片處理區 =============

def extract_images(input_pdf: str, output_dir: str, page_spec: str | None = None) -> None:
    """
    Extract all images from a PDF file.

    Args:
        input_pdf: Source PDF path.
        output_dir: Directory where images will be saved.
        page_spec: Optional page specification string for selective extraction.

    Raises:
        ImportError: If PyMuPDF or Pillow is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        OSError: If the output directory cannot be created.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法提取圖片。")
    if Image is None:
        raise ImportError("Pillow 尚未安裝，無法提取圖片。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法提取圖片：{input_pdf}") from exc

    try:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        document.close()
        raise OSError(f"無法建立輸出資料夾：{output_dir}") from exc

    try:
        total_pages = document.page_count

        if page_spec is None:
            page_indexes = list(range(total_pages))
        else:
            page_indexes = parse_page_spec(page_spec, total_pages)
            if not page_indexes:
                raise ValueError("頁碼範圍解析結果為空，請確認輸入。")

        image_count = 0
        print(f"提取圖片（共 {len(page_indexes)} 頁）...")

        for page_index in tqdm(page_indexes, desc="提取圖片", unit="頁"):
            page = document[page_index]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                image_filename = output_path / f"page_{page_index + 1:03d}_img_{img_index + 1:03d}.{image_ext}"
                image_filename.write_bytes(image_bytes)
                image_count += 1

        print(f"✓ 成功提取 {image_count} 張圖片")
    finally:
        document.close()


def pdf_to_images(input_pdf: str, output_dir: str, page_spec: str | None = None, dpi: int = 300, image_format: str = "png") -> None:
    """
    Convert PDF pages to image files.

    Args:
        input_pdf: Source PDF path.
        output_dir: Directory where images will be saved.
        page_spec: Optional page specification string for selective conversion.
        dpi: Resolution for the output images (default 300).
        image_format: Output image format - 'png', 'jpg', or 'jpeg' (default 'png').

    Raises:
        ImportError: If PyMuPDF or Pillow is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If parameters are invalid.
        OSError: If the output directory cannot be created.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法轉換為圖片。")
    if Image is None:
        raise ImportError("Pillow 尚未安裝，無法轉換為圖片。")

    if dpi < 72 or dpi > 600:
        raise ValueError("DPI 應介於 72 與 600 之間。")

    image_format = image_format.lower()
    if image_format not in ("png", "jpg", "jpeg"):
        raise ValueError("圖片格式僅支援 png, jpg, jpeg。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法轉換為圖片：{input_pdf}") from exc

    try:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        document.close()
        raise OSError(f"無法建立輸出資料夾：{output_dir}") from exc

    try:
        total_pages = document.page_count

        if page_spec is None:
            page_indexes = list(range(total_pages))
        else:
            page_indexes = parse_page_spec(page_spec, total_pages)
            if not page_indexes:
                raise ValueError("頁碼範圍解析結果為空，請確認輸入。")

        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)

        print(f"轉換 PDF 為圖片（共 {len(page_indexes)} 頁，{dpi} DPI）...")

        for page_index in tqdm(page_indexes, desc="轉換圖片", unit="頁"):
            page = document[page_index]
            pix = page.get_pixmap(matrix=matrix)

            output_file = output_path / f"page_{page_index + 1:03d}.{image_format}"
            pix.save(str(output_file))

        print(f"✓ 成功轉換 {len(page_indexes)} 頁為圖片")
    finally:
        document.close()


# ============= 頁面重排區 =============

def reorder_pages(input_pdf: str, output_pdf: str, page_order: str) -> None:
    """
    Reorder pages in a PDF according to the specified sequence.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        page_order: Comma-separated page numbers in desired order (e.g., "3,1,2,4-6").

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the page order specification is invalid.
        OSError: If the output file cannot be written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法重排頁面。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法重排頁面：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        page_indexes = parse_page_spec(page_order, total_pages)

        if not page_indexes:
            raise ValueError("頁碼順序不可為空。")

        print(f"重排 PDF 頁面（{len(page_indexes)} 頁）...")

        # Create a new document with pages in the specified order
        new_document = fitz.open()
        for page_index in tqdm(page_indexes, desc="重排頁面", unit="頁"):
            new_document.insert_pdf(document, from_page=page_index, to_page=page_index)

        try:
            new_document.save(output_pdf)
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc
        finally:
            new_document.close()

        print(f"✓ 成功重排 {len(page_indexes)} 頁")
    finally:
        document.close()


# ============= 頁首頁尾區 =============

def add_page_numbers(
    input_pdf: str,
    output_pdf: str,
    position: str = "bottom-right",
    format_str: str = "{page}",
    size: int = 10,
    color: tuple[float, float, float] = (0.0, 0.0, 0.0),
    offset: int = 20,
) -> None:
    """
    Add page numbers to all pages in a PDF.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        position: Position of page numbers - 'top-left', 'top-center', 'top-right',
                  'bottom-left', 'bottom-center', 'bottom-right' (default 'bottom-right').
        format_str: Format string for page numbers, use {page} for current page (default '{page}').
        size: Font size for page numbers (default 10).
        color: Text color as RGB tuple with values in [0, 1] (default black).
        offset: Offset in points from the edge (default 20).

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If parameters are invalid.
        OSError: If the output file cannot be written.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，無法添加頁碼。")

    valid_positions = {
        "top-left", "top-center", "top-right",
        "bottom-left", "bottom-center", "bottom-right"
    }
    if position not in valid_positions:
        raise ValueError(f"位置必須為以下之一：{', '.join(valid_positions)}")

    if "{page}" not in format_str:
        raise ValueError("格式字串必須包含 {{page}} 佔位符。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法添加頁碼：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        print(f"添加頁碼到 {total_pages} 頁...")

        for page_index in tqdm(range(total_pages), desc="添加頁碼", unit="頁"):
            page = document[page_index]
            page_num_text = format_str.format(page=page_index + 1)

            # Calculate position
            page_rect = page.rect
            text_width = fitz.get_text_length(page_num_text, fontsize=size)

            if "left" in position:
                x = offset
            elif "right" in position:
                x = page_rect.width - text_width - offset
            else:  # center
                x = (page_rect.width - text_width) / 2

            if "top" in position:
                y = offset + size
            else:  # bottom
                y = page_rect.height - offset

            point = fitz.Point(x, y)
            page.insert_text(
                point,
                page_num_text,
                fontsize=size,
                color=color,
            )

        try:
            document.save(output_pdf)
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc

        print(f"✓ 成功添加頁碼到 {total_pages} 頁")
    finally:
        document.close()


# ============= 壓縮優化區 =============


def _downsample_images_once(pdf_path: Path, scale_factor: float, jpeg_quality: int) -> bool:
    """
    Perform a single pass of image downsampling on the given PDF.

    Args:
        pdf_path: Path to the PDF file to process (overwritten in place).
        scale_factor: Factor (0-1] used to resize image dimensions.
        jpeg_quality: JPEG quality (1-95) for recompressed images.

    Returns:
        True if any images were downsampled, False otherwise.
    """
    if fitz is None or Image is None:
        return False

    doc = fitz.open(pdf_path.as_posix())
    updated = False

    try:
        for page_index in range(len(doc)):
            image_entries = doc.get_page_images(page_index, full=True)
            if not image_entries:
                continue

            page = doc.load_page(page_index)
            for entry in image_entries:
                xref = entry[0]
                pix = fitz.Pixmap(doc, xref)

                # Skip extremely small images or masks
                if pix.width < 32 or pix.height < 32 or pix.n == 0:
                    continue

                if pix.n >= 5:  # Convert CMYK/others to RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)

                mode = "RGB" if pix.n >= 3 else "L"
                image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)

                new_width = max(1, int(pix.width * scale_factor))
                new_height = max(1, int(pix.height * scale_factor))
                if new_width == pix.width and new_height == pix.height and jpeg_quality >= 80:
                    continue  # Nothing to change

                if scale_factor < 0.99:
                    image = image.resize((new_width, new_height), Image.LANCZOS)

                buffer = io.BytesIO()
                # Always save as RGB JPEG to maximize compression
                jpeg_image = image.convert("RGB")
                jpeg_image.save(buffer, format="JPEG", quality=jpeg_quality, optimize=True)
                buffer.seek(0)

                page.replace_image(xref, stream=buffer.getvalue())
                updated = True
    finally:
        if updated:
            temp_path = pdf_path.with_suffix(".tmp_optim.pdf")
            doc.save(temp_path.as_posix(), garbage=4, deflate=True, clean=True)
            doc.close()
            temp_path.replace(pdf_path)
        else:
            doc.close()

    return updated


def _apply_low_quality_downsampling(pdf_path: Path, target_ratio: float) -> None:
    """
    Aggressively downsample images to approach a desired size reduction.

    Args:
        pdf_path: Path to the PDF file to optimize.
        target_ratio: Target fraction of the original file size (e.g., 0.5 for 50%).
    """
    if fitz is None or Image is None:
        print("⚠ 低品質壓縮需要 PyMuPDF 與 Pillow，已改為保留基礎壓縮。")
        return

    if target_ratio <= 0 or target_ratio >= 1:
        target_ratio = 0.5

    original_size = pdf_path.stat().st_size
    if original_size <= 0:
        return

    target_size = int(original_size * target_ratio)
    current_size = original_size
    scale_factor = 0.78
    jpeg_quality = 62
    passes = 0

    while current_size > target_size and passes < 3:
        passes += 1
        updated = _downsample_images_once(pdf_path, scale_factor, jpeg_quality)
        if not updated:
            break
        current_size = pdf_path.stat().st_size
        scale_factor = max(0.55, scale_factor * 0.85)
        jpeg_quality = max(38, jpeg_quality - 8)

def optimize_pdf(
    input_pdf: str,
    output_pdf: str,
    linearize: bool = False,
    aggressive: bool = False,
    dpi: int = 150,
    quality_level: str | None = None,
    remove_unused: bool = True,
    compress_images: bool = True,
    remove_duplicates: bool = True,
    target_reduction: float | None = None,
) -> None:
    """
    Optimize a PDF file using pikepdf, optionally enabling linearization or aggressive image handling.

    Args:
        input_pdf: Source PDF path.
        output_pdf: Destination PDF path.
        linearize: Enable Fast-Web-View linearization when True.
        aggressive: Placeholder flag for advanced image resampling (not fully implemented).
        dpi: Target DPI for aggressive mode (validated but not applied in simplified workflow).

    Raises:
        ImportError: If required libraries are not installed.
        FileNotFoundError: If the input PDF does not exist.
        PermissionError: If the input PDF is encrypted.
        ValueError: If DPI is outside the accepted range.
        OSError: If writing the output file fails.
    """
    if pikepdf is None:
        raise ImportError("pikepdf 尚未安裝，無法執行壓縮優化。")
    if aggressive and (fitz is None or Image is None):
        print("⚠ 進階壓縮需要 PyMuPDF 與 Pillow，改為執行基礎壓縮。")
        aggressive = False
    if not 72 <= dpi <= 300:
        raise ValueError("DPI 參數應介於 72 與 300 之間。")

    source_path = Path(input_pdf)
    if not source_path.is_file():
        raise FileNotFoundError(f"找不到檔案：{source_path.resolve()}")

    original_size = source_path.stat().st_size
    mode_label = "進階" if aggressive else "基礎"
    print(f"優化 PDF（{mode_label}模式）...")

    try:
        with pikepdf.open(input_pdf) as pdf:
            if remove_unused and hasattr(pdf, "remove_unreferenced_resources"):
                try:
                    pdf.remove_unreferenced_resources()
                except Exception:
                    pass

            if remove_duplicates and hasattr(pdf, "remove_duplicate_font_dicts"):
                try:
                    pdf.remove_duplicate_font_dicts()
                except Exception:
                    pass

            for _ in tqdm(["結構清理"], desc="壓縮 PDF", unit="步驟"):
                pdf.save(
                    output_pdf,
                    linearize=linearize,
                    object_stream_mode=pikepdf.ObjectStreamMode.generate,
                    compress_streams=True,
                    stream_decode_level=pikepdf.StreamDecodeLevel.generalized,
                )
    except FileNotFoundError:
        raise
    except pikepdf.PasswordError as exc:  # type: ignore[attr-defined]
        raise PermissionError(f"檔案已加密，無法壓縮：{input_pdf}") from exc
    except pikepdf.PdfError as exc:  # type: ignore[attr-defined]
        raise ValueError(f"無法讀取 PDF 檔案：{input_pdf}") from exc
    except OSError as exc:
        raise OSError(f"無法寫入輸出檔案：{output_pdf}") from exc

    output_path = Path(output_pdf)

    quality = (quality_level or "").lower()
    desired_ratio = target_reduction if target_reduction else None

    if desired_ratio is not None:
        desired_ratio = max(0.05, min(0.95, desired_ratio))

    if quality == "low" and compress_images:
        _apply_low_quality_downsampling(output_path, desired_ratio or 0.5)
    elif aggressive and compress_images:
        _apply_low_quality_downsampling(output_path, desired_ratio or 0.6)

    new_size = output_path.stat().st_size

    saved_ratio = 0.0
    if original_size > 0:
        saved_ratio = max(0.0, (1 - new_size / original_size) * 100)

    print(
        f"✓ 壓縮完成：{format_file_size(original_size)} → {format_file_size(new_size)}"
        f"（節省 {saved_ratio:.1f}%）"
    )

    if aggressive:
        print("⚠ 簡化版進階壓縮已執行基礎壓縮，圖片重採樣功能待後續補強。")


# ============= OCR 文字識別區 =============

def extract_text_from_pdf_ocr(
    input_pdf: str,
    language: str = "eng",
    dpi: int = 300,
) -> str:
    """
    Extract text from a PDF using OCR (Optical Character Recognition).

    This function is useful for scanned PDFs or image-based PDFs where text
    cannot be extracted directly. It converts each page to an image and uses
    Tesseract OCR to recognize the text.

    Args:
        input_pdf: Source PDF path.
        language: OCR language code (default "eng" for English).
                  Common codes: eng, chi_sim, chi_tra, fra, deu, spa, jpn, etc.
        dpi: DPI resolution for rendering pages (default 300).

    Returns:
        Extracted text content from all pages.

    Raises:
        ImportError: If required libraries (PyMuPDF, Pillow, pytesseract) are not installed.
        FileNotFoundError: If the input PDF does not exist or Tesseract is not installed.
        PermissionError: If the input PDF is encrypted.
        ValueError: If the PDF cannot be read.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) 尚未安裝，請先執行 'pip install PyMuPDF>=1.23.0'。")
    if Image is None:
        raise ImportError("Pillow 尚未安裝，請先執行 'pip install Pillow>=10.0.0'。")
    if pytesseract is None:
        raise ImportError("pytesseract 尚未安裝，請先執行 'pip install pytesseract>=0.3.10'。")

    try:
        document = safe_open_pdf(input_pdf)
    except PermissionError as exc:
        raise PermissionError(f"檔案已加密，無法進行 OCR：{input_pdf}") from exc

    try:
        total_pages = document.page_count
        print(f"正在進行 OCR 文字識別（共 {total_pages} 頁，語言：{language}）...")

        extracted_text = []

        for page_index in tqdm(range(total_pages), desc="OCR 識別", unit="頁"):
            page = document[page_index]

            # Render page as image with specified DPI
            matrix = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=matrix)

            # Convert PyMuPDF pixmap to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            # Perform OCR on the image
            try:
                page_text = pytesseract.image_to_string(img, lang=language)
                if page_text.strip():
                    extracted_text.append(f"--- 第 {page_index + 1} 頁 ---\n")
                    extracted_text.append(page_text)
                    extracted_text.append("\n")
            except pytesseract.TesseractNotFoundError as exc:
                raise FileNotFoundError(
                    "Tesseract OCR 引擎未安裝。請先安裝 Tesseract：\n"
                    "  - Ubuntu/Debian: sudo apt-get install tesseract-ocr\n"
                    "  - macOS: brew install tesseract\n"
                    "  - Windows: 從 https://github.com/UB-Mannheim/tesseract/wiki 下載安裝"
                ) from exc
    finally:
        document.close()

    result = "".join(extracted_text)
    print(f"✓ OCR 識別完成，共擷取 {len(result)} 個字元")
    return result


def save_text_to_docx(text: str, output_path: str) -> None:
    """
    Save extracted text to a Microsoft Word (.docx) document.

    Args:
        text: Text content to save.
        output_path: Destination .docx file path.

    Raises:
        ImportError: If python-docx is not installed.
        OSError: If the output file cannot be written.
    """
    if Document is None:
        raise ImportError(
            "python-docx 尚未安裝，請先執行 'pip install python-docx>=0.8.11'。"
        )

    doc = Document()
    doc.add_heading("OCR 擷取文字", level=1)

    # Split text into paragraphs and add to document
    paragraphs = text.split("\n")
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)

    try:
        output_file = Path(output_path)
        output_parent = output_file.parent
        if output_parent and not output_parent.exists():
            output_parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
    except OSError as exc:
        raise OSError(f"無法寫入輸出檔案：{output_path}") from exc

    print(f"✓ 已儲存為 Microsoft Word 格式：{output_path}")


def save_text_to_odt(text: str, output_path: str) -> None:
    """
    Save extracted text to a LibreOffice Writer (.odt) document.

    Args:
        text: Text content to save.
        output_path: Destination .odt file path.

    Raises:
        ImportError: If odfpy is not installed.
        OSError: If the output file cannot be written.
    """
    if not odf_available:
        raise ImportError(
            "odfpy 尚未安裝，請先執行 'pip install odfpy>=1.4.1'。"
        )

    doc = OpenDocumentText()

    # Add a title
    title_style = Style(name="Title", family="paragraph")
    title_style.addElement(TextProperties(fontsize="18pt", fontweight="bold"))
    doc.styles.addElement(title_style)

    title = P(text="OCR 擷取文字", stylename=title_style)
    doc.text.addElement(title)

    # Add paragraphs
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            para = P(text=para_text)
            doc.text.addElement(para)

    try:
        output_file = Path(output_path)
        output_parent = output_file.parent
        if output_parent and not output_parent.exists():
            output_parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
    except OSError as exc:
        raise OSError(f"無法寫入輸出檔案：{output_path}") from exc

    print(f"✓ 已儲存為 LibreOffice Writer 格式：{output_path}")


def ocr_pdf_to_text(
    input_pdf: str,
    output_docx: str | None = None,
    output_odt: str | None = None,
    output_txt: str | None = None,
    language: str = "eng",
    dpi: int = 300,
) -> str:
    """
    Perform OCR on a PDF and save the extracted text to various formats.

    Args:
        input_pdf: Source PDF path.
        output_docx: Optional path to save as Microsoft Word (.docx).
        output_odt: Optional path to save as LibreOffice Writer (.odt).
        output_txt: Optional path to save as plain text (.txt).
        language: OCR language code (default "eng").
        dpi: DPI resolution for rendering pages (default 300).

    Returns:
        Extracted text content.

    Raises:
        ValueError: If no output format is specified.
        Same exceptions as extract_text_from_pdf_ocr, save_text_to_docx, and save_text_to_odt.
    """
    if not output_docx and not output_odt and not output_txt:
        raise ValueError("請至少指定一個輸出格式（--docx、--odt 或 --txt）。")

    # Extract text using OCR
    text = extract_text_from_pdf_ocr(input_pdf, language=language, dpi=dpi)

    # Save to requested formats
    if output_docx:
        save_text_to_docx(text, output_docx)

    if output_odt:
        save_text_to_odt(text, output_odt)

    if output_txt:
        try:
            output_file = Path(output_txt)
            output_parent = output_file.parent
            if output_parent and not output_parent.exists():
                output_parent.mkdir(parents=True, exist_ok=True)
            output_file.write_text(text, encoding="utf-8")
            print(f"✓ 已儲存為純文字格式：{output_txt}")
        except OSError as exc:
            raise OSError(f"無法寫入輸出檔案：{output_txt}") from exc

    return text


def build_parser() -> "argparse.ArgumentParser":
    """Construct the CLI argument parser."""
    parser = argparse.ArgumentParser(
        prog="pdf_toolkit",
        description="Python PDF 工具箱 - 提供合併、拆分、編輯、壓縮等功能",
        epilog="範例：python pdf_toolkit.py merge file1.pdf file2.pdf -o output.pdf",
    )
    parser.add_argument("--version", action="version", version=f"PDF Toolkit {__version__}")

    subparsers = parser.add_subparsers(dest="command", required=True, help="可用的子命令")

    merge_parser = subparsers.add_parser("merge", help="合併多個 PDF 檔案")
    merge_parser.add_argument("inputs", nargs="+", help="輸入 PDF 檔案（可多個）")
    merge_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")

    split_parser = subparsers.add_parser("split", help="拆分 PDF 檔案")
    split_parser.add_argument("input", help="輸入 PDF 檔案")
    split_parser.add_argument("-d", "--dir", dest="directory", required=True, help="輸出資料夾")
    split_parser.add_argument("-p", "--pages", help="頁碼範圍（未提供則拆分為單頁）")

    delete_parser = subparsers.add_parser("delete", help="刪除指定頁面")
    delete_parser.add_argument("input", help="輸入 PDF 檔案")
    delete_parser.add_argument("-p", "--pages", required=True, help="要刪除的頁碼範圍")
    delete_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")

    rotate_parser = subparsers.add_parser("rotate", help="旋轉指定頁面")
    rotate_parser.add_argument("input", help="輸入 PDF 檔案")
    rotate_parser.add_argument("-p", "--pages", required=True, help="要旋轉的頁碼範圍")
    rotate_parser.add_argument(
        "-a",
        "--angle",
        type=int,
        required=True,
        choices=[90, 180, 270],
        help="旋轉角度（90/180/270）",
    )
    rotate_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")

    watermark_parser = subparsers.add_parser("watermark", help="為 PDF 添加文字水印")
    watermark_parser.add_argument("input", help="輸入 PDF 檔案")
    watermark_parser.add_argument("-t", "--text", required=True, help="水印文字")
    watermark_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    watermark_parser.add_argument("--size", type=int, default=36, help="字體大小（預設 36）")
    watermark_parser.add_argument("--alpha", type=float, default=0.3, help="透明度 0-1（預設 0.3）")
    watermark_parser.add_argument("--angle", type=int, default=0, choices=[0, 90, 180, 270], help="旋轉角度，僅支援 0/90/180/270（預設 0）")

    optimize_parser = subparsers.add_parser("optimize", help="壓縮優化 PDF 檔案")
    optimize_parser.add_argument("input", help="輸入 PDF 檔案")
    optimize_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    optimize_parser.add_argument("--linearize", action="store_true", help="啟用線性化（Fast-Web-View）")
    optimize_parser.add_argument("--aggressive", action="store_true", help="啟用進階壓縮（圖片重採樣）")
    optimize_parser.add_argument("--dpi", type=int, default=150, help="進階模式的目標 DPI（預設 150）")

    info_parser = subparsers.add_parser("info", help="查詢 PDF 資訊")
    info_parser.add_argument("input", help="輸入 PDF 檔案")

    autofill_parser = subparsers.add_parser("autofill", help="自動填寫 PDF 表單")
    autofill_parser.add_argument("template", help="PDF 範本檔案")
    autofill_parser.add_argument("-o", "--output", help="輸出 PDF 檔案")
    autofill_parser.add_argument("-d", "--data", help="JSON 資料檔案路徑")
    autofill_parser.add_argument(
        "-v",
        "--value",
        action="append",
        dest="values",
        metavar="KEY=VALUE",
        help="直接指定欄位值，可重複使用（格式：key=value）",
    )
    autofill_parser.add_argument(
        "--list-fields",
        action="store_true",
        help="列出 PDF 中可填寫的表單欄位",
    )
    autofill_parser.add_argument(
        "--flatten",
        action="store_true",
        help="填寫完成後將表單欄位壓平成一般文字",
    )

    extract_parser = subparsers.add_parser("extract-text", help="從 PDF 提取文字")
    extract_parser.add_argument("input", help="輸入 PDF 檔案")
    extract_parser.add_argument("-o", "--output", help="輸出文字檔案（未指定則輸出到螢幕）")
    extract_parser.add_argument("-p", "--pages", help="要提取的頁碼範圍")

    encrypt_parser = subparsers.add_parser("encrypt", help="加密 PDF 檔案")
    encrypt_parser.add_argument("input", help="輸入 PDF 檔案")
    encrypt_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    encrypt_parser.add_argument("-u", "--user-password", required=True, help="使用者密碼（開啟檔案需要）")
    encrypt_parser.add_argument("--owner-password", help="擁有者密碼（權限管理，預設同使用者密碼）")

    decrypt_parser = subparsers.add_parser("decrypt", help="解密 PDF 檔案")
    decrypt_parser.add_argument("input", help="輸入 PDF 檔案")
    decrypt_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    decrypt_parser.add_argument("-p", "--password", required=True, help="PDF 密碼")

    extract_images_parser = subparsers.add_parser("extract-images", help="從 PDF 提取圖片")
    extract_images_parser.add_argument("input", help="輸入 PDF 檔案")
    extract_images_parser.add_argument("-d", "--dir", dest="directory", required=True, help="輸出資料夾")
    extract_images_parser.add_argument("-p", "--pages", help="要提取的頁碼範圍")

    pdf2img_parser = subparsers.add_parser("pdf-to-images", help="將 PDF 頁面轉換為圖片")
    pdf2img_parser.add_argument("input", help="輸入 PDF 檔案")
    pdf2img_parser.add_argument("-d", "--dir", dest="directory", required=True, help="輸出資料夾")
    pdf2img_parser.add_argument("-p", "--pages", help="要轉換的頁碼範圍")
    pdf2img_parser.add_argument("--dpi", type=int, default=300, help="圖片解析度（預設 300）")
    pdf2img_parser.add_argument("--format", default="png", choices=["png", "jpg", "jpeg"], help="圖片格式（預設 png）")

    reorder_parser = subparsers.add_parser("reorder", help="重排 PDF 頁面順序")
    reorder_parser.add_argument("input", help="輸入 PDF 檔案")
    reorder_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    reorder_parser.add_argument("-p", "--pages", required=True, help="新的頁碼順序（例如：3,1,2,4-6）")

    page_numbers_parser = subparsers.add_parser("page-numbers", help="添加頁碼")
    page_numbers_parser.add_argument("input", help="輸入 PDF 檔案")
    page_numbers_parser.add_argument("-o", "--output", required=True, help="輸出 PDF 檔案")
    page_numbers_parser.add_argument("--position", default="bottom-right",
                                      choices=["top-left", "top-center", "top-right",
                                               "bottom-left", "bottom-center", "bottom-right"],
                                      help="頁碼位置（預設 bottom-right）")
    page_numbers_parser.add_argument("--format", default="{page}", help="頁碼格式（使用 {page} 表示頁碼，預設 '{page}'）")
    page_numbers_parser.add_argument("--size", type=int, default=10, help="字體大小（預設 10）")
    page_numbers_parser.add_argument("--offset", type=int, default=20, help="距離邊緣的偏移量（預設 20）")

    ocr_parser = subparsers.add_parser("ocr", help="使用 OCR 識別 PDF 中的文字並匯出")
    ocr_parser.add_argument("input", help="輸入 PDF 檔案")
    ocr_parser.add_argument(
        "--docx",
        help="輸出為 Microsoft Word 格式 (.docx)",
    )
    ocr_parser.add_argument(
        "--odt",
        help="輸出為 LibreOffice Writer 格式 (.odt)",
    )
    ocr_parser.add_argument(
        "--txt",
        help="輸出為純文字格式 (.txt)",
    )
    ocr_parser.add_argument(
        "--language",
        default="eng",
        help="OCR 語言代碼（預設 eng）。常用：eng, chi_sim, chi_tra, fra, deu, spa, jpn 等",
    )
    ocr_parser.add_argument(
        "--dpi",
        type=int,
        default=300,
        help="頁面渲染 DPI 解析度（預設 300）",
    )

    return parser


def main(argv: Sequence[str] | None = None) -> None:
    """Entry point for the PDF toolkit CLI."""
    parser = build_parser()
    args = parser.parse_args(argv)

    try:
        if args.command == "merge":
            merge_pdfs(args.inputs, args.output)
        elif args.command == "split":
            split_pdf(args.input, args.directory, args.pages)
        elif args.command == "delete":
            delete_pages(args.input, args.output, args.pages)
        elif args.command == "rotate":
            rotate_pages(args.input, args.output, args.pages, args.angle)
        elif args.command == "watermark":
            add_watermark(
                args.input,
                args.output,
                args.text,
                size=args.size,
                alpha=args.alpha,
                angle=args.angle,
            )
        elif args.command == "optimize":
            optimize_pdf(
                args.input,
                args.output,
                linearize=args.linearize,
                aggressive=args.aggressive,
                dpi=args.dpi,
            )
        elif args.command == "info":
            print_pdf_info(args.input)
        elif args.command == "autofill":
            if not args.list_fields and not args.output:
                raise ValueError("請指定輸出檔案 (--output) 或使用 --list-fields 查看欄位。")

            if args.list_fields:
                fields = extract_pdf_form_fields(args.template)
                if not fields:
                    print("⚠ 未偵測到任何可填寫的表單欄位。")
                else:
                    print("📋 表單欄位清單：")
                    for field in fields:
                        base = f"  • {field['name']} ({field['type']}) - 第 {field['page']} 頁"
                        value = field.get("value")
                        if value not in (None, ""):
                            base += f"，目前值：{value}"
                        print(base)
                        options = field.get("options")
                        if options:
                            print(f"      選項：{', '.join(map(str, options))}")

            if args.output:
                payload: dict[str, Any] = {}
                payload.update(_load_json_data(args.data))
                payload.update(_parse_key_value_pairs(args.values))
                filled = fill_pdf_form(
                    args.template,
                    args.output,
                    payload,
                    flatten=args.flatten,
                )
                print(
                    f"✓ 已填寫 {len(filled)} 個欄位，輸出檔案：{Path(args.output).resolve()}"
                )
        elif args.command == "extract-text":
            extract_text(args.input, args.output, args.pages)
        elif args.command == "encrypt":
            encrypt_pdf(args.input, args.output, args.user_password, args.owner_password)
        elif args.command == "decrypt":
            decrypt_pdf(args.input, args.output, args.password)
        elif args.command == "extract-images":
            extract_images(args.input, args.directory, args.pages)
        elif args.command == "pdf-to-images":
            pdf_to_images(args.input, args.directory, args.pages, args.dpi, args.format)
        elif args.command == "reorder":
            reorder_pages(args.input, args.output, args.pages)
        elif args.command == "page-numbers":
            add_page_numbers(
                args.input,
                args.output,
                position=args.position,
                format_str=args.format,
                size=args.size,
                offset=args.offset,
            )
        elif args.command == "ocr":
            ocr_pdf_to_text(
                args.input,
                output_docx=args.docx,
                output_odt=args.odt,
                output_txt=args.txt,
                language=args.language,
                dpi=args.dpi,
            )
        else:  # pragma: no cover - subparser enforces valid commands
            parser.print_help()
    except FileNotFoundError as err:
        print(f"❌ 錯誤：{err}", file=sys.stderr)
        sys.exit(1)
    except PermissionError as err:
        print(f"❌ 權限錯誤：{err}", file=sys.stderr)
        sys.exit(1)
    except (ValueError, ImportError, OSError) as err:
        print(f"❌ 錯誤：{err}", file=sys.stderr)
        sys.exit(1)
    except KeyboardInterrupt:
        print("\n⚠ 操作已取消", file=sys.stderr)
        sys.exit(130)
    except Exception as err:  # pragma: no cover - safety net
        print(f"❌ 未預期的錯誤：{err}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
